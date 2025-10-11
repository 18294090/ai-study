from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph

async def generate_pdf(paper, file_path):
    """生成PDF格式试卷"""
    doc = SimpleDocTemplate(file_path)
    styles = getSampleStyleSheet()
    elements = []
    
    # 添加标题
    elements.append(Paragraph(paper.title, styles['Title']))
    
    # 添加题目
    for i, question in enumerate(paper.questions, 1):
        elements.append(Paragraph(f"{i}. {question.content}", styles['Normal']))
        
    doc.build(elements)

async def generate_word(paper, file_path):
    """生成Word格式试卷"""
    doc = Document()
    doc.add_heading(paper.title, 0)
    
    for i, question in enumerate(paper.questions, 1):
        doc.add_paragraph(f"{i}. {question.content}")
        
    doc.save(file_path)

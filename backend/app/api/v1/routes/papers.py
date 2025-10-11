from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from app.db.session import get_db
from app.models.paper import Paper
from app.models.question import Question
from app.core.auth import get_current_user
import random
from docx import Document
from reportlab.pdfgen import canvas
import os

router = APIRouter()

@router.post("/generate")
async def generate_paper(
    paper_config: dict,
    db: AsyncSession = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """自动生成试卷"""
    # 根据配置选择题目
    questions = []
    total_score = 0
    
    for difficulty, count in paper_config["difficulty_distribution"].items():
        query = select(Question).filter(
            Question.subject_id == paper_config["subject_id"],
            Question.difficulty == int(difficulty)
        )
        result = await db.execute(query)
        available_questions = result.scalars().all()
        
        if len(available_questions) < count:
            raise HTTPException(
                status_code=400,
                detail=f"Not enough questions with difficulty {difficulty}"
            )
            
        selected = random.sample(available_questions, count)
        questions.extend(selected)
        total_score += sum(q.score for q in selected)
    
    # 创建试卷
    paper = Paper(
        title=paper_config["title"],
        creator_id=current_user.id,
        subject_id=paper_config["subject_id"],
        total_score=total_score,
        question_config=paper_config
    )
    paper.questions = questions
    
    db.add(paper)
    await db.commit()
    return paper

@router.get("/{paper_id}/export")
async def export_paper(
    paper_id: int,
    format: str = "pdf",
    db: AsyncSession = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """导出试卷"""
    paper = await db.get(Paper, paper_id)
    if not paper:
        raise HTTPException(status_code=404, detail="Paper not found")
    
    if format == "pdf":
        file_path = f"temp/{paper_id}.pdf"
        await generate_pdf(paper, file_path)
    elif format == "docx":
        file_path = f"temp/{paper_id}.docx"
        await generate_word(paper, file_path)
    else:
        raise HTTPException(status_code=400, detail="Unsupported format")
    
    return FileResponse(file_path, filename=f"{paper.title}.{format}")

async def generate_pdf(paper: Paper, file_path: str):
    """生成PDF试卷"""
    # 使用reportlab生成PDF文件
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(file_path, pagesize=letter)
    c.drawString(100, 750, f"试卷标题: {paper.title}")
    c.drawString(100, 735, f"总分: {paper.total_score}")
    c.drawString(100, 720, "题目列表:")
    
    y = 700
    for question in paper.questions:
        c.drawString(100, y, f"- {question.content} (分值: {question.score})")
        y -= 15
    
    c.save()

async def generate_word(paper: Paper, file_path: str):
    """生成Word试卷"""
    # 使用python-docx生成Word文件
    doc = Document()
    doc.add_paragraph(f"试卷标题: {paper.title}")
    doc.add_paragraph(f"总分: {paper.total_score}")
    doc.add_paragraph("题目列表:")
    
    for question in paper.questions:
        doc.add_paragraph(f"- {question.content} (分值: {question.score})")
    
    doc.save(file_path)
